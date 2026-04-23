from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_text(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── TITLE ──
doc.add_heading('Experiment Text Breakdown', 0)
add_text('This document contains all participant-facing text in both experiments.')
add_text('Constants (both apps): Group size = 5 (1 Leader + 4 Members), Endowment = 10 tokens, Multiplier = 1.8, Rounds = 10', bold=True)

add_separator()
# ═══════════════════════════════════════════════════
# EXPERIMENT 1
# ═══════════════════════════════════════════════════
doc.add_heading('EXPERIMENT 1: Fixed Leader Payment (public_goods_leader)', 1)
add_text('Leader payment: Fixed 20 tokens per round', bold=True)

# Page 1
add_separator()
doc.add_heading('Page 1: Prolific ID (Welcome to the Study)', 2)
add_quote('Thank you for participating in this research study, which will take about 10 minutes.')
add_quote('Before we begin, please enter your Prolific ID below. This is required to ensure you receive payment for your participation.')
add_quote('Your Payment: Upon completing the study, you will receive a $2.5 participation payment. In addition, there is a 25% chance you will be randomly selected to receive a bonus from the study. If you are chosen for a bonus, then points received will be converted to dollars at a rate of 25 cents per point. If selected, your bonus amount will be determined based on the procedures described later in the study and will be paid no later than December 24th.')
add_text('[Form: Prolific ID field]', italic=True)

# Page 2
add_separator()
doc.add_heading('Page 2: No-Deception Policy', 2)
add_quote('The experiment is conducted by researchers associated with the Interdisciplinary Experimental Laboratory (IELab) at Indiana University and follows strict no deception policy.')
add_quote('In accordance with IELab rules, whatever we tell you is indeed true. For example, if we tell you that something will happen if you choose a given option, then that thing will happen if you choose that option. Anything else would violate IELab policies. For more information on this policy, see https://ie.lab.indiana.edu/conducting-experiments')

# Page 3
add_separator()
doc.add_heading('Page 3: Welcome (Instructions)', 2)
add_quote('You are participating in an economic decision-making study.')
add_quote('Your earnings in this study will depend on your decisions, on the decisions of others, and on chance.')
add_quote('All decisions are anonymous. You will not learn the identity of any of the other participants, and they will not learn yours.')
add_quote('You will now proceed to the instructions. Please read them carefully.')
add_quote('You will be required to complete a short comprehension quiz before beginning the experiment.')

# Page 4
add_separator()
doc.add_heading('Page 4: Group Assignment (Instructions 2)', 2)
add_quote('You will be randomly and anonymously assigned to a group of 5 participants.')
add_quote('In each group, there are two types of participants: 1 Leader and 4 Group Members.')
add_quote('The assignment of roles is random. You will learn your role after reading the instructions.')
add_quote('All 5 participants in each group will receive the same instructions.')

# Page 5
add_separator()
doc.add_heading('Page 5: Payment Information (Instructions 3)', 2)
add_quote('At the end of the experiment:')
add_quote('Tokens earned across all rounds will be converted into real money.')
add_quote('1 token = ?')
add_quote('Your total payment equals:')
add_bullet('Your earnings from all rounds')
add_bullet('Plus a guaranteed participation payment of ?')

# Page 6
add_separator()
doc.add_heading('Page 6: Group Members (Instructions 4)', 2)
add_quote('Each group member receives an endowment of 10 tokens.')
add_quote('Each group member must decide how many tokens to allocate to the Group Account and how many tokens to keep in their Private Account.')
add_quote('The number of tokens must be a whole number between 0 and 10.')
add_quote('All decisions are made in private.')

# Page 7
add_separator()
doc.add_heading('Page 7: Accounts (Instructions 5)', 2)
add_text('Group Account', bold=True)
add_quote('Tokens contributed to the Group Account are added together to form a total amount. This total is multiplied by 1.8 and then evenly split among all 4 group members.')
add_text('Private Account', bold=True)
add_quote('Each member keeps the tokens remaining in their Private Account after contributing to the Group Account. These tokens are not multiplied.')

# Page 8
add_separator()
doc.add_heading('Page 8: Leader (Instructions 6)', 2)
add_quote("The leader's task is to send a recommendation at the start of each round, suggesting how many tokens each group member should contribute to the group account.")
add_quote('Group members are free to follow or ignore this recommendation.')
add_quote('The leader will not make a contributing decision to the group account and instead will receive a payment of half of the total contributed group account.')
add_quote('Leader payments do not come out of the group account.')

# Page 9
add_separator()
doc.add_heading('Page 9: Examples', 2)
add_quote('The following examples illustrate how payoffs are calculated.')

doc.add_heading('Leader Payoffs', 3)
add_quote('Regardless of how many tokens group members contribute to the Group Account, the Leader always receives 20 tokens.')

doc.add_heading('Group Member Payoffs', 3)
add_text('Example 1:', bold=True)
add_quote('You allocate 4 tokens and the other participants allocate 3, 5, and 8 tokens.')
add_quote('Total contribution: 4 + 3 + 5 + 8 = 20 tokens')
add_quote('20 × 1.8 = 36 tokens, so each member receives 36 / 4 = 9 tokens')
add_quote('You keep 6 tokens, so your total payoff is: 6 + 9 = 15 tokens')

add_text('Example 2:', bold=True)
add_quote('You allocate 10 tokens and the other participants allocate 0, 0, and 0 tokens.')
add_quote('Total contribution: 10 + 0 + 0 + 0 = 10 tokens')
add_quote('10 × 1.8 = 18 tokens, so each member receives 18 / 4 = 4.5 tokens')
add_quote('You keep 0 tokens, so your total payoff is: 0 + 4.5 = 4.5 tokens')

add_text('Example 3:', bold=True)
add_quote('You allocate 0 tokens and the other participants allocate 6, 6, and 6 tokens.')
add_quote('Total contribution: 0 + 6 + 6 + 6 = 18 tokens')
add_quote('18 × 1.8 = 32.4 tokens, so each member receives 32.4 / 4 = 8.1 tokens')
add_quote('You keep 10 tokens, so your total payoff is: 10 + 8.1 = 18.1 tokens')

# Page 10
add_separator()
doc.add_heading('Page 10: Comprehension Questions', 2)
add_quote('Please answer the following comprehension questions. You will not be able to continue to the experiment until all answers are correct. The instructions are below the questions as a reminder.')

add_text('Questions:', bold=True)
add_quote('1. Each token a group member contributes to the Group Account reduces their Private Account by how many tokens?')
add_text('   Correct answer: 1 | Hint: "Incorrect. Each contributed token reduces the private account by 1."', italic=True)

add_quote('2. If all 4 group members contribute 10 tokens each, how many tokens does each member earn?')
add_text('   Correct answer: 18 | Hint: "Incorrect. Multiply total contributions by 1.8 and divide equally."', italic=True)

add_quote('3. If all group members contribute 10 tokens each to the Group Account:')
add_bullet('a. Different group members would receive different earnings.')
add_bullet("b. Each group member's earnings would be lower than if no one contributed.")
add_bullet("c. Each group member's earnings would be higher than if no one contributed. ✓")
add_text('   Hint: "Incorrect. Contributions increase earnings for everyone."', italic=True)

add_quote("4. The leader's payoff in this experiment consists of:")
add_bullet('a. A fixed payment only. ✓')
add_bullet('b. A fixed payment and a performance-based payment.')
add_text('   Hint: "Incorrect. The leader receives only a fixed payment."', italic=True)

add_text('Note: First 2 wrong attempts show generic "Incorrect. Please try again." After 2 wrong attempts, specific hints are shown.', italic=True)

add_text('Instructions are repeated below the questions as a reminder (Group Assignment, Group Members, Accounts, Leader, Payment sections).', italic=True)

# Page 11
add_separator()
doc.add_heading('Page 11: Welcome to the Public Goods Game (Introduction)', 2)
add_text('If Leader:', bold=True)
add_quote('You are the Leader of your group.')
add_bullet('You will send a message to motivate the 4 group members')
add_bullet('You will receive a fixed payment of 20 tokens')
add_bullet('You do NOT receive a share of the group pot')

add_text('If Member:', bold=True)
add_quote('You are a Group Member.')
add_bullet('You have an endowment of 10 tokens')
add_bullet('You will receive a message from your Leader')
add_bullet('You will decide how much to contribute to the group account (0 to 10 tokens)')
add_bullet('All contributions will be multiplied by 1.8 and redistributed equally among all 4 group members')
add_bullet("You keep any tokens you don't contribute, plus your share of the redistributed pot")

add_quote("Ready to begin? When you click the button below, you'll wait for all other participants to be ready. Once everyone is ready, the round will begin and timers will start.")
add_text('[Button: "I\'m Ready to Start"]', italic=True)

# ── ROUND PAGES ──
add_separator()
doc.add_heading('ROUND PAGES (repeat for rounds 1–10)', 1)
add_text('Each round page shows: "Round X of 10"', italic=True)

# Round Start Wait
doc.add_heading('Round Start Wait Page', 2)
add_quote('Please Wait — Waiting for all players to be ready...')
add_text('Timer text: "Time left for group members to join:"', italic=True)

# Leader Message
add_separator()
doc.add_heading('Leader Message (Leader only)', 2)
add_quote('You have 2 minutes to write a message to motivate your group members to contribute.')
add_quote('You must submit a message within this time. Otherwise, you will be removed from the experiment.')
add_text('[Form: Message text area]', italic=True)
add_text('If members have dropped out:', bold=True)
add_quote('⚠️ X group member(s) have left the game. Your group has been reduced from 4 to Y members.')
add_text('Timer text: "Time left for the leader to write their message:"', italic=True)

# Wait for Leader
add_separator()
doc.add_heading('Waiting for the Leader (Members only)', 2)
add_quote('Please wait while the leader writes their message to the group...')
add_text('Timer text: "Time left for the leader to write their message:"', italic=True)

# Contribute
add_separator()
doc.add_heading('Message from Leader and Your Contribution Decision (Members only)', 2)
add_quote("Message from your Leader: [leader's message displayed]")
add_quote('You have 10 tokens. How much would you like to contribute to the group account?')
add_quote('Remember: All contributions will be multiplied by 1.8 and shared equally among the X active group members.')
add_text('[Form: Contribution (0–10), "✔ I\'m still here" checkbox]', italic=True)
add_text('If leader has dropped out:', bold=True)
add_quote('⚠️ Your leader has left the game. This was their last message:')
add_text('If members have dropped out:', bold=True)
add_quote('⚠️ X group member(s) have left the game. Your group has been reduced from 4 to Y members.')
add_text('Timer text: "Time left for group members to contribute:"', italic=True)

# Wait for Contributions
add_separator()
doc.add_heading('Waiting for Contributions', 2)
add_quote('Please wait while all group members submit their contributions...')
add_text('Timer text: "Time left for group members to contribute:"', italic=True)

# Results
add_separator()
doc.add_heading('Results', 2)
add_text('Group Results', bold=True)
add_quote('Total group contribution: X tokens (from Y active members)')
add_quote('Multiplied pot: Z')
add_quote('Division: Z ÷ Y active members = W per member')

add_text('If Leader:', bold=True)
add_quote('As the Leader, you receive a fixed payment of 20 points.')

add_text('If Member:', bold=True)
add_quote('Share per member: W points')
add_quote('Your contribution: X tokens')
add_quote('Tokens kept: Y tokens')
add_quote('Your share of pot: W points')

add_quote('Your total payoff: P')

add_text('If members have dropped out:', bold=True)
add_quote('⚠️ X member(s) did not contribute (left the game). Your group was reduced from 4 to Y active members. The multiplied pot is divided only among the Y active members.')
add_text('Timer text: "Time left to view results (page will automatically continue to trust ratings):"', italic=True)

# Trust - Leader
add_separator()
doc.add_heading('Trust Ratings — Leader', 2)
add_quote('Please rate your level of trust in the group members.')
add_quote('You have 2 minutes to complete this page. If you do not confirm that you are still here, you will be exited from the study.')
add_quote('On a scale from 0 (No trust at all) to 10 (Complete trust):')
add_text('[Form: "How much do you trust the other Members in this group?" (0–10)]', italic=True)
add_text('[Form: "✔ I\'m still here"]', italic=True)

# Trust - Member
add_separator()
doc.add_heading('Trust Ratings — Member', 2)
add_quote('Please rate your level of trust in your group.')
add_quote('You have 2 minutes to complete this page. If you do not confirm that you are still here, you will be exited from the study.')
add_quote('On a scale from 0 (No trust at all) to 10 (Complete trust):')
add_text('[Form: "How much do you trust the Leader in this group?" (0–10)]', italic=True)
add_text('[Form: "How much do you trust the other Members in this group?" (0–10)]', italic=True)
add_text('[Form: "✔ I\'m still here"]', italic=True)
add_text('If leader dropped out:', bold=True)
add_quote('⚠️ Your leader has left the game. The group has no leader now.')
add_text('If members dropped out:', bold=True)
add_quote('⚠️ X group member(s) left the game. Your group now has Y member(s) remaining.')
add_text('Timer text: "Time left for trust ratings to be submitted:"', italic=True)

# Trust Wait
add_separator()
doc.add_heading('Trust Ratings Wait Page', 2)
add_quote('Please wait while all group members submit their trust ratings...')
add_text('Timer text: "Time left for trust ratings to be submitted:"', italic=True)

# ── END PAGES ──
add_separator()
doc.add_heading('END OF EXPERIMENT PAGES', 1)

doc.add_heading('Exit Notice (shown if player drops out or group ends early)', 2)
add_text('If dropout:', bold=True)
add_quote('You have left this experiment.')
add_quote('Your response was not completed in time, so your participation has ended for this study.')
add_quote('If this was not expected, please contact the researcher.')

add_text('If group ended (too many dropouts):', bold=True)
add_quote('This experiment has ended.')
add_quote('Three group members have left the experiment, so the session cannot continue.')
add_quote('You will not continue to the next round.')

add_text('Both cases show:', bold=True)
add_quote('Your completion code is: CAHCFER6')
add_quote('[Click Here to Complete on Prolific] → https://app.prolific.co/submissions/complete?cc=CAHCFER6')

add_separator()
doc.add_heading('Completion (shown after round 10 for active players)', 2)
add_quote('Thank You!')
add_quote('Confirm Your Completion — To receive payment on Prolific, please click the button below to confirm your completion.')
add_quote('Your completion code is: CAHCFER6')
add_quote('[Click Here to Complete on Prolific] → https://app.prolific.co/submissions/complete?cc=CAHCFER6')
add_quote("If the button doesn't work, please copy and paste this URL into your browser: (URL shown)")

# ═══════════════════════════════════════════════════
# EXPERIMENT 2
# ═══════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('EXPERIMENT 2: Variable Leader Payment (public_goods_leader_variable)', 1)
add_text('Leader payment: 50% of total contributions (variable)', bold=True)
add_text('All pages are identical to Experiment 1 EXCEPT for the following differences:', italic=True)

add_separator()
doc.add_heading('Differences in Instructions 6 (Leader)', 2)
add_text('Same as Experiment 1 — both say:', italic=True)
add_quote('The leader will not make a contributing decision to the group account and instead will receive a payment of half of the total contributed group account.')
add_quote('Leader payments do not come out of the group account.')

add_separator()
doc.add_heading('Differences in Examples Page', 2)
add_text('Leader Payoffs (Variable version):', bold=True)
add_quote('The Leader receives half of the total contributions to the Group Account. This payment does not come out of the Group Account.')
add_quote('If total contributions = 20 tokens, the Leader receives 20 / 2 = 10 tokens.')
add_text('(Experiment 1 says: "the Leader always receives 20 tokens")', italic=True)
add_text('Group member payoff examples are identical.', italic=True)

add_separator()
doc.add_heading('Differences in Comprehension Questions', 2)
add_text('Questions 1–3 are identical. Question 4 differs:', italic=True)

add_text('Experiment 1 — Question 4:', bold=True)
add_quote("The leader's payoff in this experiment consists of:")
add_bullet('a. A fixed payment only. ✓')
add_bullet('b. A fixed payment and a performance-based payment.')

add_text('Experiment 2 — Question 4:', bold=True)
add_quote("The leader's payoff in this experiment consists of:")
add_bullet('a. 50% of total contributions. ✓')
add_bullet('b. 50% of the multiplied pot.')
add_bullet('c. A fixed payment only.')

add_separator()
doc.add_heading('Differences in Comprehension Page — Leader Reminder', 2)
add_text('Experiment 1:', bold=True)
add_quote('The leader does not contribute and instead receives a fixed payment of 20 tokens. This payment does not depend on group contributions.')
add_text('Experiment 2:', bold=True)
add_quote('The leader does not contribute and instead receives 50% of the total contributions. Leader payments do not reduce group members\' earnings.')

add_separator()
doc.add_heading('Differences in Introduction (Role Reveal)', 2)
add_text('Leader intro — Experiment 1:', bold=True)
add_bullet('You will receive a fixed payment of 20 tokens')
add_bullet('You do NOT receive a share of the group pot')
add_text('Leader intro — Experiment 2:', bold=True)
add_bullet('You will receive 50% of the total contributions made by the group members')
add_bullet('You do NOT receive a share of the group pot')
add_text('Member intro is identical in both experiments.', italic=True)

add_separator()
doc.add_heading('Differences in Results Page', 2)
add_text('Leader results line — Experiment 1:', bold=True)
add_quote('As the Leader, you receive a fixed payment of 20 points.')
add_text('Leader results line — Experiment 2:', bold=True)
add_quote('As the Leader, you receive 50% of total contributions: X points.')
add_text('Member results are identical in both experiments.', italic=True)

add_separator()
doc.add_heading('All Other Pages', 2)
add_text('All other pages (ProlificID, No-Deception Policy, Welcome, Group Assignment, Payment Information, Group Members, Accounts, all wait pages, contribution page, trust ratings, exit notice, completion) are identical between the two experiments.')

# Save
doc.save('/Users/lily/Desktop/svs/experiment_text_breakdown.docx')
print('Done! Saved to experiment_text_breakdown.docx')
